from docx import Document

doc = Document('Project Description and Rubrics .docx')

with open('docx_content.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        f.write(para.text + '\n')
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            f.write('\t'.join(row_text) + '\n')
        f.write('\n')